# Author : Harith Abeysinghe
# File Name : GUI.py
# Date : 29/09/2023


import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
import docx
import os


# Function to save the text
def save_text():
    global current_file
    text = text_entry.get("1.0", "end-1c")  # Get the text from the text entry widget
    if text:
        if current_file:
            if current_file.endswith(".txt"):  # Check if the current file has a .txt extension
                with open(current_file, "w") as file:
                    file.write(text)  # Save the text to a .txt file
            elif current_file.endswith(".docx"):  # Check if the current file has a .docx extension
                doc = docx.Document()
                doc.add_paragraph(text)
                doc.save(current_file)  # Save the text as a .docx file
            messagebox.showinfo("Info", f"Saved to {current_file}")  # Show a message box with a success message
        else:
            save_as()  # If there is no current file, prompt the user to save with a new file name


# Function to open a file
def open_file():
    global current_file
    file_name = filedialog.askopenfilename(defaultextension=".docx",
                                           filetypes=[("Text Documents", "*.txt"), ("Word Documents", "*.docx")])
    if file_name:
        current_file = file_name
        file_extension = os.path.splitext(file_name)[1]  # Get the file extension
        if file_extension == ".txt":  # If the file is a .txt file
            with open(file_name, "r") as file:
                text = file.read()  # Read the text from the file
        elif file_extension == ".docx":  # If the file is a .docx file
            doc = docx.Document(file_name)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])  # Read the text from the document
        text_entry.delete("1.0", "end")  # Clear the current text in the text entry widget
        text_entry.insert("1.0", text)  # Insert the loaded text into the text entry widget


# Function to save as a new file
def save_as():
    global current_file
    file_name = filedialog.asksaveasfilename(defaultextension=".txt",
                                             filetypes=[("Text Documents", "*.txt"), ("Word Documents", "*.docx")])
    if file_name:
        current_file = file_name
        save_text()  # Call the save_text function to save the text with the new file name


# Create the main window
root = tk.Tk()
root.title("Text Editor")
root.geometry("700x900")

# Create buttons for Open, Save, and Save As
open_button = tk.Button(root, text="Open", command=open_file)
save_button = tk.Button(root, text="Save", command=save_text)
save_as_button = tk.Button(root, text="Save As", command=save_as)

# Create a text entry widget
text_entry = tk.Text(root, wrap=tk.WORD, width=80, height=30)

# Grid configuration for buttons and text entry
open_button.grid(row=0, column=0, padx=10, pady=10, sticky="w")
save_button.grid(row=0, column=1, padx=10, pady=10, sticky="w")
save_as_button.grid(row=0, column=2, padx=10, pady=10, sticky="w")
text_entry.grid(row=1, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

# Initialize the current file variable
current_file = None

# Configure grid weights for resizing
root.grid_rowconfigure(1, weight=1)
root.grid_columnconfigure(0, weight=1)

# Run the GUI application
root.mainloop()
